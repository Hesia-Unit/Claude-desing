#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ACCENT = RGBColor(0x0F, 0x3B, 0x66)
ACCENT_SOFT = RGBColor(0x2B, 0x6C, 0x91)
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x58, 0x63, 0x72)
CODE_BG = "EEF2F7"
RULE_BG = "D7E3F0"


@dataclass
class CoverMetadata:
    title: str
    subtitle: str
    language: str
    fields: List[Tuple[str, str]]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def ensure_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def configure_document(document: Document, doc_title: str) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.8)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title_style = document.styles["Title"]
    title_style.font.name = "Aptos Display"
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = ACCENT

    subtitle = ensure_style(document, "Hesia Subtitle")
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_after = Pt(6)

    h1 = document.styles["Heading 1"]
    h1.font.name = "Aptos Display"
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = ACCENT
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = document.styles["Heading 2"]
    h2.font.name = "Aptos Display"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = ACCENT_SOFT
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = document.styles["Heading 3"]
    h3.font.name = "Aptos"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = TEXT
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    code_style = ensure_style(document, "Hesia Code")
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9.2)
    code_style.font.color.rgb = RGBColor(0x16, 0x24, 0x33)
    code_style.paragraph_format.left_indent = Cm(0.5)
    code_style.paragraph_format.right_indent = Cm(0.2)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.line_spacing = 1.05

    note_style = ensure_style(document, "Hesia Note")
    note_style.font.name = "Aptos"
    note_style.font.size = Pt(9.5)
    note_style.font.italic = True
    note_style.font.color.rgb = MUTED
    note_style.paragraph_format.left_indent = Cm(0.3)
    note_style.paragraph_format.space_after = Pt(6)

    bullet_style = ensure_style(document, "Hesia Bullet")
    bullet_style.base_style = normal
    bullet_style.paragraph_format.left_indent = Cm(0.63)
    bullet_style.paragraph_format.first_line_indent = Cm(-0.35)
    bullet_style.paragraph_format.space_after = Pt(2)

    document.core_properties.title = doc_title
    document.core_properties.subject = "HESIA engineering handover"
    document.core_properties.category = "Engineering documentation"
    document.core_properties.comments = "Generated from repository markdown by render_markdown_docx.py"

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run(f"{doc_title} | ")
    footer_run.font.name = "Aptos"
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = MUTED
    page_run = footer.add_run()
    page_run.font.name = "Aptos"
    page_run.font.size = Pt(8.5)
    page_run.font.color.rgb = MUTED
    add_field(page_run, "PAGE")


def parse_cover_metadata(lines: Sequence[str], fallback_title: str, fallback_language: str) -> CoverMetadata:
    subtitle = ""
    language = fallback_language
    fields: List[Tuple[str, str]] = []
    in_cover = False
    for line in lines:
        if line.startswith("## "):
            section_name = line[3:].strip().lower()
            if section_name in {"page de garde", "cover page"}:
                in_cover = True
                continue
            if in_cover:
                break
        if not in_cover:
            continue
        match = re.match(r"\*\*(.+?)\*\*:\s*(.+)$", line.strip())
        if match:
            key = match.group(1).strip()
            value = match.group(2).strip()
            fields.append((key, value))
            lowered = key.lower()
            if lowered in {"langue", "language"}:
                language = value
            if lowered in {"scope", "portee"}:
                subtitle = value
    if not subtitle:
        subtitle = "Complete engineering handover reference"
    return CoverMetadata(title=fallback_title, subtitle=subtitle, language=language, fields=fields)


def split_body(lines: Sequence[str]) -> List[str]:
    skip_section = None
    body: List[str] = []
    in_body = False

    for line in lines:
        if line.startswith("## "):
            heading = line[3:].strip().lower()
            if heading in {"page de garde", "cover page"}:
                skip_section = heading
                continue
            if heading in {"sommaire", "table of contents"}:
                skip_section = heading
                continue
            skip_section = None
            if re.match(r"^\d+\.", heading):
                in_body = True
        if skip_section is not None:
            continue
        if in_body:
            body.append(line)
    return body


INLINE_TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_inline_runs(paragraph, text: str) -> None:
    parts = INLINE_TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.2)
            run.font.color.rgb = RGBColor(0x10, 0x3A, 0x5A)
            continue
        paragraph.add_run(part)


def add_separator(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(" ")
    p_par_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE_BG)
    p_bdr.append(bottom)
    p_par_pr.append(p_bdr)
    r.font.size = Pt(1)


def add_code_block(document: Document, lines: Sequence[str], language: str) -> None:
    if language:
        note = document.add_paragraph(style="Hesia Note")
        note.add_run(f"Code / source block ({language})")
    for line in lines:
        p = document.add_paragraph(style="Hesia Code")
        p.paragraph_format.keep_together = True
        p.paragraph_format.keep_with_next = False
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9.2)
        tc_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), CODE_BG)
        tc_pr.append(shd)


def flush_paragraph_buffer(document: Document, buffer: List[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer if part.strip())
    if text:
        p = document.add_paragraph()
        add_inline_runs(p, text)
    buffer.clear()


def add_cover(document: Document, meta: CoverMetadata) -> None:
    hero = document.add_paragraph()
    hero.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hero.paragraph_format.space_before = Pt(100)
    hero.paragraph_format.space_after = Pt(8)
    hero.style = document.styles["Title"]
    hero.add_run(meta.title)

    sub = document.add_paragraph(style="Hesia Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(18)
    sub.add_run(meta.subtitle)

    lang = document.add_paragraph(style="Hesia Subtitle")
    lang.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lang.add_run(f"Language: {meta.language}")

    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for key, value in meta.fields:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value
        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row[0], "E8F0F7")
        for run in row[0].paragraphs[0].runs:
            run.bold = True
        for cell in row:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)

    document.add_page_break()


def add_toc(document: Document, toc_title: str) -> None:
    p = document.add_paragraph(toc_title, style="Heading 1")
    p.paragraph_format.space_after = Pt(8)
    toc_p = document.add_paragraph()
    add_field(toc_p.add_run(), r'TOC \o "1-3" \h \z \u')
    note = document.add_paragraph(style="Hesia Note")
    note.add_run("Update the table of contents in Word if page numbers are not shown yet.")
    document.add_page_break()


def render_body(document: Document, lines: Sequence[str]) -> None:
    paragraph_buffer: List[str] = []
    in_code = False
    code_lang = ""
    code_lines: List[str] = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                flush_paragraph_buffer(document, paragraph_buffer)
                add_code_block(document, code_lines, code_lang)
                code_lines.clear()
                code_lang = ""
                in_code = False
            else:
                code_lines.append(line)
            continue

        if stripped.startswith("```"):
            flush_paragraph_buffer(document, paragraph_buffer)
            in_code = True
            code_lang = stripped[3:].strip()
            continue

        if not stripped:
            flush_paragraph_buffer(document, paragraph_buffer)
            continue

        if stripped == "---":
            flush_paragraph_buffer(document, paragraph_buffer)
            add_separator(document)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph_buffer(document, paragraph_buffer)
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            style = {2: "Heading 1", 3: "Heading 2", 4: "Heading 3"}.get(level, "Heading 1")
            document.add_paragraph(title, style=style)
            continue

        bullet_match = re.match(r"^-\s+(.+)$", stripped)
        if bullet_match:
            flush_paragraph_buffer(document, paragraph_buffer)
            p = document.add_paragraph(style="Hesia Bullet")
            p.style = "List Bullet"
            add_inline_runs(p, bullet_match.group(1))
            continue

        num_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_match:
            flush_paragraph_buffer(document, paragraph_buffer)
            p = document.add_paragraph(style="List Number")
            add_inline_runs(p, num_match.group(1))
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph_buffer(document, paragraph_buffer)
    if in_code:
        add_code_block(document, code_lines, code_lang)


def render_markdown_to_docx(input_path: Path, output_path: Path, title: str, language_hint: str) -> None:
    raw_text = input_path.read_text(encoding="utf-8")
    lines = raw_text.splitlines()

    meta = parse_cover_metadata(lines, fallback_title=title, fallback_language=language_hint)
    body = split_body(lines)

    document = Document()
    configure_document(document, meta.title)
    add_cover(document, meta)
    add_toc(document, "Sommaire" if meta.language.lower().startswith("fr") else "Table of Contents")
    render_body(document, body)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a repository markdown reference into a styled DOCX.")
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--title", default="")
    parser.add_argument("--language", default="English")
    args = parser.parse_args()

    title = args.title.strip() or args.input.stem.replace("_", " ")
    render_markdown_to_docx(args.input, args.output, title=title, language_hint=args.language)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
